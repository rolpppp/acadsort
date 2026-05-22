"""DOCX text extraction using python-docx."""

import logging
from pathlib import Path

from docx import Document

from backend.config import settings
from backend.extraction.utils import clean_and_truncate, normalize_text

logger = logging.getLogger(__name__)


async def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted and cleaned text, truncated to max chars
        
    Raises:
        ValueError: If file cannot be read or is corrupted
    """
    try:
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise ValueError(f"DOCX file not found: {file_path}")
        
        if file_path.stat().st_size > settings.max_file_size_mb * 1024 * 1024:
            raise ValueError(
                f"DOCX file too large: {file_path.stat().st_size / 1024 / 1024:.1f}MB "
                f"(max: {settings.max_file_size_mb}MB)"
            )
        
        logger.info(f"Extracting text from DOCX: {file_path}")
        
        doc = Document(file_path)
        text_parts = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = "\n".join(text_parts)
        
        if not full_text.strip():
            raise ValueError("No text could be extracted from DOCX")
        
        # Normalize and truncate
        normalized = normalize_text(full_text)
        cleaned = clean_and_truncate(normalized, settings.extraction_max_chars)
        
        logger.info(f"DOCX extraction complete: {len(cleaned)} chars from {len(text_parts)} elements")
        
        return cleaned
        
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {e}")
        raise


async def extract_text_from_docx_with_metadata(file_path: str) -> dict:
    """
    Extract text and metadata from a DOCX file.
    
    Returns dict with keys:
        - text: Extracted text
        - element_count: Number of paragraphs + table cells
        - title: Document title (if available)
    """
    try:
        file_path = Path(file_path)
        doc = Document(file_path)
        
        # Extract metadata
        core_properties = doc.core_properties
        
        text_parts = []
        element_count = 0
        
        # Paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                element_count += 1
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
                        element_count += 1
        
        full_text = "\n".join(text_parts)
        normalized = normalize_text(full_text)
        cleaned = clean_and_truncate(normalized, settings.extraction_max_chars)
        
        return {
            "text": cleaned,
            "element_count": element_count,
            "title": core_properties.title or "",
            "author": core_properties.author or "",
        }
        
    except Exception as e:
        logger.error(f"DOCX metadata extraction failed for {file_path}: {e}")
        raise
